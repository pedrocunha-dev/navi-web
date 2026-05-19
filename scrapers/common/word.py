import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

IMG_WIDTH  = Cm(7.0)
IMG_HEIGHT = Cm(5.0)


def _criar_documento() -> Document:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)

    cabecalho = doc.add_paragraph()
    cabecalho.alignment = 1
    run = cabecalho.add_run(
        "[Organization]\n[Department]\n"
        "PHOTO REPORT"
    )
    run.bold = True
    doc.add_paragraph("")
    return doc


def _listar_imagens(pasta_imovel: str, max_imgs: int) -> list:
    return sorted([
        os.path.join(pasta_imovel, f)
        for f in os.listdir(pasta_imovel)
        if f.lower().endswith(('.jpg', '.jpeg', '.png'))
    ])[:max_imgs]


def _inserir_cabecalho_imovel(doc: Document, id_imovel: str, lote: str, link: str) -> None:
    p = doc.add_paragraph()
    p.add_run("ID: ").bold = True
    p.add_run(id_imovel)

    p = doc.add_paragraph()
    p.add_run("Batch: ").bold = True
    p.add_run(lote)

    p = doc.add_paragraph()
    p.add_run("Link: ").bold = True
    p.add_run(link)


def _inserir_tabela_2x3_merged(doc: Document, imagens: list) -> None:
    """3x2 table with last row merged — layout for Site 2 / Site 3."""
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for row in range(3):
        if row < 2:
            for col in range(2):
                cell = table.cell(row, col)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if idx < len(imagens):
                    try:
                        cell.paragraphs[0].add_run().add_picture(imagens[idx], width=IMG_WIDTH, height=IMG_HEIGHT)
                    except Exception as e:
                        print(f"Error inserting image {imagens[idx]}: {e}")
                    idx += 1
        else:
            merged = table.cell(row, 0).merge(table.cell(row, 1))
            merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx < len(imagens):
                try:
                    p = merged.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(imagens[idx], width=IMG_WIDTH, height=IMG_HEIGHT)
                except Exception as e:
                    print(f"Error inserting image {imagens[idx]}: {e}")
                idx += 1


def _inserir_tabela_2x4(doc: Document, imagens: list) -> None:
    """4x2 table without merging — layout for Site 1."""
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for row in range(4):
        for col in range(2):
            cell = table.cell(row, col)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx < len(imagens):
                try:
                    cell.paragraphs[0].add_run().add_picture(imagens[idx], width=IMG_WIDTH, height=IMG_HEIGHT)
                except Exception as e:
                    print(f"Error inserting image {imagens[idx]}: {e}")
                idx += 1


def salvar_relatorio_shared(data: list, nome_docx: str, lote_numero: str) -> None:
    """Generates the photo report using the 3x2 merged-row layout (Site 2 / Site 3)."""
    doc = _criar_documento()
    for entry in data:
        _inserir_cabecalho_imovel(doc, entry["ID"], lote_numero, entry["Link"])
        doc.add_paragraph("")
        imagens = _listar_imagens(entry["Caminho na Pasta"], 5)
        _inserir_tabela_2x3_merged(doc, imagens)
        doc.add_page_break()
    doc.save(nome_docx)
    print(f"Word document '{nome_docx}' saved successfully.")


def salvar_relatorio_site1(data: list, nome_docx: str, lote_numero: str, raspar_descricao: bool = False) -> None:
    """Generates the photo report for Site 1.

    - Normal mode (raspar_descricao=False): 4x2 table, up to 8 images.
    - Description mode (raspar_descricao=True): 3x2 merged-row table, up to 5 images.
    """
    doc = _criar_documento()
    max_fotos = 5 if raspar_descricao else 8
    inserir_tabela = _inserir_tabela_2x3_merged if raspar_descricao else _inserir_tabela_2x4
    for entry in data:
        _inserir_cabecalho_imovel(doc, entry["ID"], lote_numero, entry["Link"])
        imagens = _listar_imagens(entry["Caminho na Pasta"], max_fotos)
        inserir_tabela(doc, imagens)
        doc.add_paragraph("")
        doc.add_paragraph("")
    doc.save(nome_docx)
    print(f"Word document '{nome_docx}' saved successfully.")
