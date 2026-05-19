import os
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

from scrapers.common.word import (
    _inserir_tabela_2x3_merged,
    _inserir_tabela_2x4,
)


def gerar_relatorio_fotografico():
    pasta_base = Path.home() / "AppData" / "Local" / "NAVI" / "tmp" / "imobiliario"
    pasta_saida = Path.home() / "Downloads"
    arquivo_excel = pasta_base / "planilha_geral_preliminar.xlsx"
    documento_saida = pasta_saida / "relatorio_fotografico_consolidado.docx"

    df = pd.read_excel(arquivo_excel)

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)

    header = document.add_paragraph()
    header.alignment = 1
    run = header.add_run(
        "[Organization]\n[Department]\n"
        "CONSOLIDATED PHOTO REPORT"
    )
    run.bold = True
    document.add_paragraph("")

    for _, row in df.iterrows():
        numero = row['Nº']
        id_imovel = row['ID']
        lote = row['Lote']
        fonte = row['Fonte']
        link = row['Link']

        fonte_pasta = fonte.lower().replace(" ", "_")
        pasta_imagens = pasta_base / f"imagens_{fonte_pasta}_lote_{lote}" / f"imovel_{id_imovel}"

        if not pasta_imagens.exists():
            continue

        imagens_todas = sorted([
            str(pasta_imagens / img) for img in os.listdir(pasta_imagens)
            if img.lower().endswith(('.jpg', '.jpeg', '.png'))
        ])

        if not imagens_todas:
            continue

        # Site 1 in normal mode downloads up to 8 photos -> 4x2 layout.
        # Site 1 in description mode downloads up to 5 photos -> 3x2 merged layout.
        # Site 2 / Site 3 always download up to 5 photos -> 3x2 merged layout.
        # Detection is based on the actual photo count on disk.
        usar_layout_site1_4x2 = fonte.lower() == "site1" and len(imagens_todas) > 5

        if usar_layout_site1_4x2:
            imagens = imagens_todas[:8]
        else:
            imagens = imagens_todas[:5]

        paragrafo = document.add_paragraph()
        run = paragrafo.add_run(f"Nº {numero} - Link: ")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run2 = paragrafo.add_run(link)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(12)

        if usar_layout_site1_4x2:
            _inserir_tabela_2x4(document, imagens)
        else:
            document.add_paragraph("")
            _inserir_tabela_2x3_merged(document, imagens)

        document.add_page_break()

    document.save(documento_saida)
    print(f"Consolidated photo report saved to: {documento_saida}")

    shutil.rmtree(pasta_base)
